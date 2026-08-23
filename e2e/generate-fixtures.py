#!/usr/bin/env python3
"""Generate the E2E fixture corpus: fictional agency documents in every format
DocSage ingests, plus two rejection cases. Deterministic (seeded), pure-fake
content — no real people, agencies, or data.

Runs inside the backend uv environment (python-docx, openpyxl, pillow,
reportlab are dev dependencies there):

    uv run python e2e/generate-fixtures.py --out e2e/fixtures
"""

from __future__ import annotations

import argparse
import csv
import io
import random
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

random.seed(20260828)
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"


def _font(size: int, bold: bool = False):
    try:
        return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)
    except OSError:
        return ImageFont.load_default()


def make_org_chart_png(path: Path) -> None:
    img = Image.new("RGB", (980, 620), (250, 250, 248))
    d = ImageDraw.Draw(img)
    d.text((40, 32), "Meridian County Office of Resilience — Reporting Structure",
           font=_font(24, True), fill=(28, 32, 40))
    boxes = [
        ((390, 90, 590, 150), "Director V. Okafor"),
        ((180, 230, 380, 290), "Preparedness Lead S. Tanaka"),
        ((600, 230, 800, 290), "Recovery Lead A. Whitfield"),
        ((60, 380, 260, 440), "Planner J. Ruiz"),
        ((300, 380, 500, 440), "Trainer K. Mensah"),
        ((540, 380, 740, 440), "Grants Officer P. Lindqvist"),
        ((780, 380, 950, 440), "Analyst D. Petrov"),
    ]
    links = [(0, 1), (0, 2), (1, 3), (1, 4), (2, 5), (2, 6)]
    for a, b in links:
        x1, y1, x2, y2 = boxes[a][0][2], boxes[a][0][3], boxes[b][0][0], boxes[b][0][1]
        cx1, cy1 = boxes[a][0][0] + (boxes[a][0][2] - boxes[a][0][0]) // 2, boxes[a][0][3]
        cx2, cy2 = boxes[b][0][0] + (boxes[b][0][2] - boxes[b][0][0]) // 2, boxes[b][0][1]
        d.line((cx1, cy1, cx1, (cy1 + cy2) // 2, cx2, (cy1 + cy2) // 2, cx2, cy2), fill=(120, 128, 140), width=3)
    for (x0, y0, x1, y1), label in boxes:
        d.rounded_rectangle((x0, y0, x1, y1), 10, fill=(255, 255, 255), outline=(70, 110, 160), width=2)
        d.text(((x0 + x1) // 2 - int(len(label) * 3.4), (y0 + y1) // 2 - 10), label,
               font=_font(13), fill=(28, 32, 40))
    d.text((40, 540), "Chart 4.1 — Incident Command alignment, adopted March 2026 review",
           font=_font(15), fill=(96, 102, 112))
    img.save(path)


def make_site_photo_jpg(path: Path) -> None:
    img = Image.new("RGB", (860, 640))
    d = ImageDraw.Draw(img)
    for y in range(640):
        t = y / 640
        d.line([(0, y), (860, y)], fill=(int(180 - 60 * t), int(190 - 40 * t), int(210 - 20 * t)))
    d.polygon([(0, 500), (220, 380), (430, 500), (640, 400), (860, 520), (860, 640), (0, 640)],
              fill=(122, 128, 118))
    d.ellipse((640, 80, 740, 180), fill=(250, 214, 120))
    for x, w, h in ((120, 60, 190), (300, 46, 150), (560, 70, 220)):
        d.rectangle((x, 640 - h, x + w, 640), fill=(150, 148, 152))
        d.rectangle((x + 8, 640 - h + 10, x + w - 8, 640 - h + 34), fill=(220, 226, 232))
    d.rectangle((20, 20, 330, 52), fill=(24, 30, 42))
    d.text((30, 28), "Photo 7 — Elm Creek staging area, drill 2026-06-11",
           font=_font(15, True), fill=(240, 244, 250))
    img.save(path, quality=88)


def make_handbook_docx(path: Path, chart_png: Path) -> None:
    doc = Document()
    doc.add_heading("Volunteer Onboarding Handbook", 0)
    doc.add_paragraph("Meridian County Office of Resilience · Edition 4 · Effective April 2026")
    doc.add_heading("1. Enrollment steps", 1)
    for step in (
        "Complete the interest form at any branch office or through the county portal.",
        "Attend a 90-minute orientation session offered every second Tuesday.",
        "Pass the standard background screening; results arrive within ten business days.",
        "Select an assignment track: preparedness education, logistics support, or radio communications.",
        "Collect the badge and equipment kit from the Elm Creek facility front desk.",
    ):
        doc.add_paragraph(step, style="List Number")
    doc.add_heading("2. Assignment tracks and time commitments", 1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    rows = (
        ("Track", "Minimum hours / month", "Supervisor"),
        ("Preparedness education", "6", "S. Tanaka"),
        ("Logistics support", "10", "A. Whitfield"),
        ("Radio communications", "8", "J. Ruiz"),
    )
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.rows[i].cells[j].text = value
    doc.add_heading("3. Reporting structure", 1)
    doc.add_picture(str(chart_png), width=Inches(6.0))
    doc.add_paragraph(
        "Figure 1 — where volunteers fit in the reporting structure. Escalation "
        "questions go to the track supervisor first; the duty officer line "
        "(extension 4412) is staffed during activations only."
    )
    doc.add_heading("4. Reimbursement", 1)
    doc.add_paragraph(
        "Mileage is reimbursed at the county rate of 58.5 cents per mile for "
        "approved assignments. Submit travel claims within 30 days through the "
        "finance portal; the grants officer, P. Lindqvist, approves claims under "
        "$500 and routes larger amounts to the director."
    )
    doc.save(path)


def make_grants_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Awards"
    header_fill = PatternFill("solid", fgColor="2F5D8C")
    ws.append(["Award ID", "Funder", "Purpose", "Awarded ($)", "Spent ($)", "Remaining ($)"])
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
    awards = [
        ("RES-2026-011", "State Resilience Fund", "Cooling center upgrades", 148_000, 91_250),
        ("RES-2026-014", "Harbor Foundation", "Flood sensor mesh pilot", 82_500, 82_500),
        ("RES-2026-019", "Metro United Way", "Preparedness education kits", 36_000, 12_400),
        ("RES-2026-023", "County Bridge Fund", "Generator replacement, Elm Creek", 61_750, 0),
        ("RES-2026-027", "State Resilience Fund", "Volunteer radio refresh", 27_300, 9_850),
    ]
    for row in awards:
        ws.append(list(row))
    for i in range(2, 7):
        ws.cell(row=i, column=6, value=f"=D{i}-E{i}")
    ws.append(["TOTAL", "", "", "=SUM(D2:D6)", "=SUM(E2:E6)", "=SUM(F2:F6)"])
    for cell in ws[7]:
        cell.font = Font(bold=True)
    for column, width in zip("ABCDEF", (16, 24, 30, 13, 12, 13)):
        ws.column_dimensions[column].width = width
    notes = wb.create_sheet("Notes")
    notes.append(["Note", "Detail"])
    notes.append(["Closeout", "RES-2026-014 closeout due December 15, 2026; final report drafted by the analyst."])
    notes.append(["Match", "Harbor Foundation requires a 10% match tracked in the county bridge allocation."])
    wb.save(path)


def make_inspection_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=letter)
    w, h = letter
    c.setFont("Helvetica-Bold", 16)
    c.drawString(inch, h - 0.75 * inch, "Shelter Facility Inspection — Northgate Community Center")
    c.setFont("Helvetica", 10)
    c.drawString(inch, h - 0.98 * inch, "Inspector: K. Mensah · Visit date: June 3, 2026 · Report RES-INSP-118")
    rows = [
        ("Checklist item", "Status", "Follow-up"),
        ("Generator fuel level (diesel)", "Pass — 82%", "None"),
        ("Potable water storage (gallons)", "Pass — 1,100", "None"),
        ("Cots staged (count)", "Fail — 140 of 220", "Order 80 by July 1"),
        ("Emergency lighting test", "Pass", "None"),
        ("Kitchen refrigeration", "Fail — unit 2 at 46F", "Service call 4812"),
        ("ADA restroom access", "Pass", "None"),
    ]
    y = h - 1.5 * inch
    for i, (item, status, follow) in enumerate(rows):
        font, size = ("Helvetica-Bold", 10) if i == 0 else ("Helvetica", 10)
        c.setFont(font, size)
        c.drawString(inch, y, item)
        c.drawString(inch * 3.5, y, status)
        c.drawString(inch * 5.6, y, follow)
        y -= 0.28 * inch
    c.setFont("Helvetica", 10)
    for line in (
        "Summary: two corrective actions opened. The cot shortfall is the pacing item for",
        "hurricane-season readiness; the facilities lead confirmed delivery window June 26-28.",
        "Re-inspection scheduled for July 8, 2026 with the same checklist.",
    ):
        y -= 0.24 * inch
        c.drawString(inch, y, line)
    c.showPage()
    c.setFont("Helvetica-Bold", 13)
    c.drawString(inch, h - inch, "Appendix — photolog reference")
    c.setFont("Helvetica", 10)
    c.drawString(inch, h - 1.35 * inch,
                 "Photos 1-6 filed under RES-INSP-118 in the shared drive; see also the staging")
    c.drawString(inch, h - 1.6 * inch,
                 "area photo in the Elm Creek drill record for layout comparison.")
    c.showPage()
    c.save()


def make_incident_log_txt(path: Path) -> None:
    entries = [
        "2026-06-11 08:42 — Elm Creek flood drill activated. Radio net on channel 3.",
        "2026-06-11 09:05 — Staging area at the north lot confirmed by planner Ruiz.",
        "2026-06-11 09:30 — Simulated levee breach at marker 7; sandbag team dispatched.",
        "2026-06-11 10:15 — Transport request for two medical manikins to triage tent B.",
        "2026-06-11 11:40 — Drill paused 20 minutes for a real lost-child report; resolved.",
        "2026-06-11 13:00 — Hot wash notes: radio discipline improved; cot count short again.",
        "2026-06-11 15:10 — Drill concluded. After-action report assigned to analyst Petrov.",
    ]
    path.write_text("ELM CREEK FLOOD DRILL — INCIDENT LOG\n" + "\n".join(entries) + "\n")


def make_press_release_md(path: Path) -> None:
    path.write_text(
        "# County Opens Third Cooling Center for Summer 2026\n\n"
        "**Meridian County Office of Resilience — for immediate release**\n\n"
        "Beginning June 20, the Fairview branch library joins the cooling center\n"
        "network alongside Northgate Community Center and the Harborview YMCA.\n"
        "All three sites will run noon to 8 p.m. on days the heat index is\n"
        "forecast above 100 degrees.\n\n"
        "## What to bring\n\n"
        "- Identification is welcome but not required\n"
        "- Water is provided; containers encouraged\n"
        "- Service animals are welcome at every site\n\n"
        "## Transportation\n\n"
        "Route 14 buses will add a Fairview stop on activation days. Paratransit\n"
        "riders can book cooling-center trips without the usual 24-hour notice\n"
        "by quoting code COOL-26.\n\n"
        "Volunteers interested in staffing sign-in desks should complete the\n"
        "orientation in the volunteer handbook; the next session is the second\n"
        "Tuesday of July.\n"
    )


def make_roster_csv(path: Path) -> None:
    with path.open("w", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(["badge", "name", "track", "orientation_date", "background_check", "hours_ytd"])
        writer.writerows([
            ["V-0107", "Rowan Ashby", "Preparedness education", "2026-03-10", "clear", 41],
            ["V-0112", "Dana Whitcomb", "Logistics support", "2026-03-10", "clear", 78],
            ["V-0118", "Milo Ferreira", "Radio communications", "2026-04-14", "clear", 52],
            ["V-0121", "Priya Nandakumar", "Preparedness education", "2026-04-14", "pending", 6],
            ["V-0126", "Elias Brandt", "Logistics support", "2026-05-12", "clear", 19],
            ["V-0130", "Freya Osei", "Radio communications", "2026-05-12", "clear", 33],
        ])


def make_invalid_zip(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("readme.txt", "Zips are not an accepted document type; DocSage should refuse this upload.")


def make_oversized_txt(path: Path) -> None:
    with path.open("w") as fh:
        chunk = "Padding line for the size guard test — no meaningful content.\n" * 64
        written = 0
        while written < 26 * 1024 * 1024:
            fh.write(chunk)
            written += len(chunk)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out", required=True, type=Path, help="output directory")
    args = parser.parse_args()
    out: Path = args.out
    out.mkdir(parents=True, exist_ok=True)

    chart = out / "org-chart.png"
    make_org_chart_png(chart)
    make_site_photo_jpg(out / "site-photo.jpg")
    make_handbook_docx(out / "volunteer-handbook.docx", chart)
    make_grants_xlsx(out / "grants-tracking.xlsx")
    make_inspection_pdf(out / "shelter-inspection.pdf")
    make_incident_log_txt(out / "drill-incident-log.txt")
    make_press_release_md(out / "press-release.md")
    make_roster_csv(out / "volunteer-roster.csv")
    make_invalid_zip(out / "invalid-payload.zip")
    make_oversized_txt(out / "oversized-blob.txt")

    sizes = {p.name: p.stat().st_size for p in sorted(out.iterdir())}
    for name, size in sizes.items():
        print(f"{name:28} {size:>10,} bytes")
    print(f"{len(sizes)} fixtures written to {out}")


if __name__ == "__main__":
    main()
