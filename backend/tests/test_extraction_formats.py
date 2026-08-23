"""Golden tests for format-level extraction: the real binaries committed under
db/seed-corpus are pushed through ``extract`` directly (no HTTP, no database)
and the resulting parts are asserted against the corpus's known contents."""

import io
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image
from reportlab.pdfgen import canvas

from docsage_api.services.extraction.base import extract, markdown_table

SEED_CORPUS = Path(__file__).resolve().parents[2] / "db" / "seed-corpus"

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def corpus_copy(tmp_path: Path, name: str) -> Path:
    """Extract from a tmp copy so tests can never mutate the committed corpus."""
    destination = tmp_path / name
    shutil.copy2(SEED_CORPUS / name, destination)
    return destination


def test_docx_paragraphs_and_table_follow_document_order(tmp_path):
    result = extract(corpus_copy(tmp_path, "telework-policy.docx"), DOCX_MIME)
    contents = [p.content for p in result.parts]
    assert all(p.kind in ("text", "table") for p in result.parts)

    # The eligibility table serializes as a markdown table with header and data rows.
    table = next(p for p in result.parts if p.kind == "table")
    assert "| Employee category | Max remote days/week | Approval authority |" in table.content
    assert "| Full-time staff | 3 | Supervisor + Director |" in table.content
    assert "| Contractors | 2 | Contracting Officer Representative |" in table.content

    # Prose paragraphs survive verbatim, including the policy's key phrase.
    assert any("management option, not an employee right" in content for content in contents)

    # The five numbered headings and the table keep their document order.
    markers = [
        "1. Purpose",
        "2. Eligibility",
        "3. Maximum telework days",
        "4. Technical requirements",
        "5. Records and confidentiality",
    ]
    positions = [contents.index(marker) for marker in markers]
    assert positions == sorted(positions)
    assert positions[2] < contents.index(table.content) < positions[3]


def test_docx_embedded_image_becomes_image_part(tmp_path):
    # The corpus docx files carry no media parts, so exercise the image
    # relationship path with a document built here: python-docx embeds the
    # PNG blob verbatim, which is what the extractor must surface.
    document = Document()
    document.add_paragraph("Chart intro.")
    document.add_picture(str(SEED_CORPUS / "ticket-volume-chart.png"), width=Inches(2))
    document.add_paragraph("Chart outro.")
    path = tmp_path / "embeds-chart.docx"
    document.save(str(path))

    result = extract(Path(path), DOCX_MIME)
    images = [p for p in result.parts if p.kind == "image"]
    assert len(images) == 1
    assert images[0].image_bytes.startswith(b"\x89PNG\r\n\x1a\n")
    assert images[0].mime == "image/png"
    assert images[0].filename.endswith(".png")
    assert {p.content for p in result.parts if p.kind == "text"} == {"Chart intro.", "Chart outro."}


def test_xlsx_serializes_every_sheet_with_quarterly_values(tmp_path):
    result = extract(corpus_copy(tmp_path, "fy2027-it-budget.xlsx"), XLSX_MIME)
    assert [p.kind for p in result.parts] == ["table", "table"]

    budget, notes = (p.content for p in result.parts)
    assert "### Sheet: FY2027 Budget" in budget
    assert "FY2027" in budget
    assert "| Category | Q1 | Q2 | Q3 | Q4 | Total |" in budget
    # Quarterly figures are read from the stored cell values.
    assert "| Conference room AV | 0 | 95 | 140 | 0 |  |" in budget
    assert "| Security appliances | 0 | 0 | 275 | 0 |  |" in budget
    assert "| Software licensing | 310 | 310 | 310 | 310 |  |" in budget

    assert "### Sheet: Notes" in notes
    assert "Office of the Chief Information Officer" in notes


def test_pdf_corpus_text_is_page_addressed(tmp_path):
    result = extract(corpus_copy(tmp_path, "records-retention-schedule.pdf"), "application/pdf")
    assert result.page_count == 1
    text_parts = [p for p in result.parts if p.kind == "text"]
    assert [p.page for p in text_parts] == [1]
    body = text_parts[0].content
    assert "NARA General Records Schedule" in body
    assert "Equipment inventory" in body
    assert "Life of asset + 3 years" in body


def test_pdf_multi_page_documents_carry_page_numbers(tmp_path):
    # The corpus PDF is single-page, so generate a two-page one to pin the
    # per-page splitting and page numbering behavior.
    path = tmp_path / "two-pages.pdf"
    writer = canvas.Canvas(str(path))
    writer.drawString(72, 720, "Northgate facility inspection checklist")
    writer.showPage()
    writer.drawString(72, 720, "Re-inspection is required within 30 days")
    writer.showPage()
    writer.save()

    result = extract(path, "application/pdf")
    assert result.page_count == 2
    text_parts = [p for p in result.parts if p.kind == "text"]
    assert [p.page for p in text_parts] == [1, 2]  # one part per page, in order
    assert "Northgate" in text_parts[0].content
    assert "Re-inspection" in text_parts[1].content


def test_png_reencodes_as_single_image_part(tmp_path):
    result = extract(corpus_copy(tmp_path, "ticket-volume-chart.png"), "image/png")
    assert len(result.parts) == 1
    part = result.parts[0]
    assert part.kind == "image"
    assert part.image_bytes.startswith(b"\x89PNG\r\n\x1a\n")
    assert part.mime == "image/png"
    assert part.filename == "ticket-volume-chart.png"
    with Image.open(io.BytesIO(part.image_bytes)) as reencoded:
        assert reencoded.size == (880, 520)  # under MAX_EDGE, dimensions preserved


def test_csv_renders_as_markdown_table(tmp_path):
    result = extract(corpus_copy(tmp_path, "equipment-inventory.csv"), "text/csv")
    assert [p.kind for p in result.parts] == ["table"]
    lines = result.parts[0].content.splitlines()
    assert lines[0] == "| asset_tag | item | purchased | assigned_to | condition |"
    assert lines[1] == "| --- | --- | --- | --- | --- |"
    assert all(line.startswith("|") and line.endswith("|") for line in lines)
    assert any("AST-004182" in line for line in lines)
    assert any("MacBook Pro 14 M4" in line for line in lines)


def test_txt_and_md_pass_through_verbatim(tmp_path):
    cases = [
        ("facilities-call-notes.txt", "text/plain"),
        ("migration-notes.md", "text/markdown"),
    ]
    for name, mime in cases:
        path = corpus_copy(tmp_path, name)
        result = extract(path, mime)
        assert len(result.parts) == 1
        assert result.parts[0].kind == "text"
        assert result.parts[0].content == path.read_text(encoding="utf-8")


def test_markdown_table_escapes_pipes_and_flattens_newlines():
    rows = [
        ["Metric", "Note"],
        ["Uptime|SLO", "line one\nline two"],
        [None, ""],
    ]
    assert markdown_table(rows) == (
        "| Metric | Note |\n"
        "| --- | --- |\n"
        "| Uptime\\|SLO | line one line two |\n"
        "|  |  |"
    )
    assert markdown_table([]) == ""
